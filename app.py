from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from datetime import datetime 
from docx2pdf import convert
from flask import Flask, request, jsonify, session
from openai import AzureOpenAI  
import os
import subprocess
import shutil
import sys
# import win32com.client
# import pythoncom
from docx import Document
import fitz
from dotenv import load_dotenv
import os
import re
load_dotenv()
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text =[]
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)
def extract_text_from_pdf(pdf_path):
    text=''
    with fitz.open(pdf_path) as pdf_doc:
        for page in pdf_doc:
            text += page.get_text()
    return text
app = Flask(__name__)
#Placeholder dict
app.secret_key = 'Vantive'
placeholder_values = {}
expected_placeholders =["Author name"]

from docx import Document

def replace_placeholders_in_docx(input_path, output_path, replacements):
    doc = Document(input_path)
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if f"{{{key}}}" in para.text:
                para.text = para.text.replace(f"{{{key}}}", val)
    doc.save(output_path)
def find_placeholders(doc_path):
    doc = Document(doc_path)
    placeholders= set()
    for para in doc.paragraphs:
        matches= re.findall(r"\{\{(.*?)\}\}", para.text)
        for match in matches:
            placeholders.add(match.strip())
    return list(placeholders)

CORS(app)
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///sessions.db"
db = SQLAlchemy(app)
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
model_deployed="gpt-4o" #"gpt-35-turbo"
client = AzureOpenAI(
    api_key=os.getenv("AZURE_API_KEY"),
    api_version=os.getenv("AZURE_API_VERSION"),
    azure_endpoint=os.getenv("AZURE_API_ENDPOINT")
)

@app.route('/')
def landing():
    return render_template('landing.html')
@app.route('/app')
def index():
    return render_template('index.html')
@app.route('/upload', methods=['POST'])

def upload_file():
    file = request.files['file']
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(file_path)
    session['uploaded_file'] = file_path
    session['awaiting_placeholder'] = False

    placeholders= find_placeholders(file_path)
    session['expected_placeholders'] = placeholders
    session['filled_placeholders'] ={}
    return jsonify({
        "message": f"File uploaded successfully. Let's start filling in the placeholders: {placeholders}.",
        "placeholders": placeholders})

#PH-NEW
    # file = request.files['file']
    # filename = secure_filename(file.filename)
    # ext = filename.rsplit('.', 1)[-1].lower()
    # filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    # file.save(filepath)

    return jsonify({
        "pdf_url": "/uploads/{}".format(filename),
        "type": "pdf",
        "filename": filename
    })
    # uploaded_file_path = filepath
    # text_content = ""
    # if ext == 'docx':
    #     pythoncom.CoInitialize()
    #     word = win32com.client.Dispatch("Word.Application")
    #     word.Visible = False
    #     pdf_filename = filename.rsplit('.',1)[0]+'.pdf'
    #     pdf_path = os.path.abspath(os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename))
    #     try:
    #         doc = word.Documents.Open(os.path.abspath(filepath))
    #         doc.ExportAsFixedFormat(pdf_path,17)
    #         doc.Close()
    #     except Exception as e:
    #         word.Quit()
    #         print(f"Error:{e}")
    #         return jsonify({"error":str(e)}),500
    #     finally:
    #         word.Quit()
    #     return jsonify({
    #         "pdf_url": "/uploads/{}".format(pdf_filename),
    #         "type": "pdf",
    #         "filename": pdf_filename
    #     })
    # elif ext == 'pdf':
    #     return jsonify({
    #         "pdf_url": f"/uploads/{filename}",
    #         "type": "pdf",
    #         "filename": filename
    #     })
    return jsonify({"error": "Unsupported file type"}), 400
@app.route('/uploads/<path:filename>')
def serve_uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)
def chunk_text(text, max_chars=3000):
    paragraphs = text.split('\n')
    chunks, current = [], ""
    for para in paragraphs:
        if len(current) + len(para) < max_chars:
            current += para + "\n"
        else:
            chunks.append(current.strip())
            current = para + "\n"
    if current:
        chunks.append(current.strip())
    return chunks

questions = [{"key":"AUTHOR", "question": "Who is the Author for this report?"},
                {"key":"SYSTEM_ENGG", "question": "Who is the System engg for this?"}]
@app.route('/chat', methods=['POST'])
def chat():
    user_input = request.json['message']
    expected = session.get('expected_placeholders', [])
    filled= session.get('filled_placeholders', {})
    awaiting = session.get('awaiting_placeholder')

    if awaiting:
        filled[awaiting]= user_input
        session['filled_placeholders']= filled
        session['awaiting_placeholder']= None
    if set(filled.keys()) == set(expected):
        replace_placeholders_in_docx(
            session['uploaded_file'],
            'output.docx',
            filled
        )
    
    # if session.get('awaiting_author'):
    #     placeholder_values['Author name'] = user_input
    #     session['awaiting_author'] = False

    #     # Replace placeholders now
    #     replace_placeholders_in_docx(
    #         session['uploaded_file'],
    #         'output.docx',
    #         placeholder_values
    #     )
        return jsonify({"response": "Thanks! The document has been updated with the author name."})
    else:
        next_placeholder = list(set(expected) - set(filled.keys()))[0]
        session['awaiting_placeholder']= next_placeholder
        return jsonify({"response": f"What is '{next_placeholder}'?"})
    if expected and set(filled.keys()) != set(expected):
        next_placeholder= list(set(expected)- set(filled.keys()))[0]
        session['awaiting_placeholder']= next_placeholder
        return jsonify({"response": f"What is '{next_placeholder}'?"})
    

    # if 'author' in user_input.lower():
    #     session['awaiting_author'] = True
    #     return jsonify({'response': 'Who is the Author for this report?'})

    return jsonify({'response': 'What would you like to do next?'})







    

    if user_input == "start":
        # reference_path = os.path.join(app.config['UPLOAD_FOLDER'], 'filled_reference.docx')
        # uploaded_filename = request.json.get('uploaded_filename')
        uploaded_filename = "template_111.docx"
        uploaded_path = os.path.join(app.config['UPLOAD_FOLDER'], uploaded_filename)
        # Extract text
        # if uploaded_filename.lower().endswith(".docx"):
        #     uploaded_text = extract_text_from_docx(uploaded_path)
        # elif uploaded_filename.lower().endswith(".pdf"):
        #     uploaded_text = extract_text_from_pdf(uploaded_path)
        # else:
        #     return jsonify({"error": "Unsupported file format"}), 400
        # reference_text = extract_text_from_docx(reference_path)
        # Chunk both documents
        # ref_chunks = chunk_text(reference_text)
        # uploaded_chunks = chunk_text(uploaded_text)
        #Test

        # print("=== Reference Text ===")
        # print(reference_text[:1000])  
        # print("=== Uploaded Text ===")
        # print(uploaded_text[:1000])

        # Generate questions
#         for ref, up in zip(ref_chunks, uploaded_chunks):
#             # if ref.strip() == up.strip():
#             #     continue #To skip similar chunks
#             prompt = f"""
# Compare the following two sections of a document. Identify what’s missing or incomplete in the uploaded version, and generate upto 2 specific, concise  questions  that woud help a user fill in the missing information.
# --- Reference ---
# {ref}
# --- Uploaded ---
# {up}
# """
#             print("=== Prompt Sent to Azure OpenAI ===")
#             print(prompt)
#             try:
#                 response = client.chat.completions.create(
#                     model=model_deployed,
#                     messages=[
#                         {"role": "system", "content": "You are a helpful assistant."},
#                         {"role": "user", "content": prompt}
#                     ]
#                 )
#                 questions_text=(response.choices[0].message.content.strip())
#                 #questions.append(questions_text)
#                 print("=== AI Response ===")
#                 print(questions_text)
#                 for q in questions_text.split("\n"):
#                     clean_q = q.strip("- ").strip()
#                     if clean_q:
#                         questions.append(clean_q)

            # except Exception as e:
            #     print(f"Error from Azure OpenAI: {e}")
            #     continue
        # Flatten and clean questions
        #all_questions = "\n".join(questions)
        return jsonify({"reply":questions})
    else:
        print("response from UI:", user_input)
        replacements = {
            "{{AUTHOR_NAME}}": "John Doe",
            "{{REASON_FOR_CHANGE}}": "Design update for regulatory compliance"
        }
        replace_placeholders(
            docx_path="uploads/final_output_doc.docx",
            output_path="uploads/final_output_doc_filled.docx",
            replacements=replacements
        )
        replace_placeholders()
        return jsonify({'reply': f"Got it! You said: {user_input}"})

def replace_placeholders(docx_path, output_path, replacements):
    doc = Document(docx_path)
    # Replace in paragraphs
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                inline = para.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(key, val)
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)
    doc.save(output_path)

@app.route('/download')
def download_filled():
    directory = os.path.join(app.root_path, 'filled_docs')
    filename = 'output_filled.pdf'
    if not os.path.exists(os.path.join(directory, filename)):
        return f"File not found at: {os.path.join(directory, filename)}", 404
    return send_from_directory(directory, filename, as_attachment=True)
if __name__ == '__main__':
    app.run(debug=True, port=5050)